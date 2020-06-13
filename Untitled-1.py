#   git add . && git commit -m "Editado por Pedro" && git push --force
#   git add . && git push --force
from docx import Document

doc = Document("Covid-19.docx")

# Adicionando pergunta teste

doc.add_paragraph("Pergunta Teste?")
doc.add_paragraph("Resposta Teste.")


doc.save("Covid-19.docx")